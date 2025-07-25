import os
import io
import keyboard
import win32clipboard
from docx import Document
from docx.shared import Inches
from PIL import ImageGrab, Image

# 🕒 STEP 1: Ask for timeframe
print("Enter the timeframe (e.g. 1m, 5s, 2m, 15m, etc.):")
timeframe = input("🕒 Timeframe: ").strip()
folder_name = f"{timeframe}_dataset"
folder_path = os.path.join("D:/trading_web_bot/dataset", folder_name)
os.makedirs(folder_path, exist_ok=True)

# 📁 Word file paths for screenshots
UP_PATH = os.path.join(folder_path, f"{timeframe}_UP.docx")
DOWN_PATH = os.path.join(folder_path, f"{timeframe}_DOWN.docx")
NO_SIGNAL_PATH = os.path.join(folder_path, f"{timeframe}_NO_SIGNAL.docx")

# 🛡 Ensure .docx files exist
for path in [UP_PATH, DOWN_PATH, NO_SIGNAL_PATH]:
    if not os.path.exists(path):
        Document().save(path)

# 🧠 Cropping area for full chart screenshot (used for overall pattern analysis)
# ➕ Increase x to move right, ➖ decrease x to move left
# ➕ Increase y to move down, ➖ decrease y to move up
# ➕ Increase width to capture more width, ➖ to reduce it
# ➕ Increase height to capture more vertical space
x_full = 150       # distance from left of screen
y_full = 220       # distance from top of screen
width_full = 530   # width of chart
height_full = 480  # height of chart

# 🧠 Cropping area for cropped 4-candle zone (advanced trim-based cropping)
# 🟩 This allows you to crop precisely FROM EACH SIDE (left/right/top/bottom)
# ➕ Increase `left_trim` to remove more from the LEFT
# ➕ Increase `right_trim` to remove more from the RIGHT
# ➕ Increase `top_trim` to remove more from the TOP
# ➕ Increase `bottom_trim` to remove more from the BOTTOM
base_x_crop = 150
base_y_crop = 220
base_width_crop = 530
base_height_crop = 480

# ✂️ Trimming amounts (you change only these to crop from edges)
left_trim = 452     # pixels to trim from LEFT
right_trim = 0      # pixels to trim from RIGHT
top_trim = 0        # pixels to trim from TOP
bottom_trim = 0     # pixels to trim from BOTTOM

# Final coordinates for cropped image
x_crop = base_x_crop + left_trim
y_crop = base_y_crop + top_trim
width_crop = base_width_crop - left_trim - right_trim
height_crop = base_height_crop - top_trim - bottom_trim

# 📋 Send image to clipboard
def send_to_clipboard(image: Image.Image):
    output = io.BytesIO()
    image.convert("RGB").save(output, "BMP")
    data = output.getvalue()[14:]
    output.close()
    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardData(win32clipboard.CF_DIB, data)
    win32clipboard.CloseClipboard()

# 📸 Capture and insert both full + cropped into 1 Word file
def capture_and_insert(doc_path, direction):
    # 🖼️ Full chart screenshot
    full_img = ImageGrab.grab(bbox=(x_full, y_full, x_full + width_full, y_full + height_full))
    full_img.save("manual_full.png")

    # 🔍 Cropped screenshot (right candles zone)
    cropped_img = ImageGrab.grab(bbox=(x_crop, y_crop, x_crop + width_crop, y_crop + height_crop))
    cropped_img.save("manual_cropped.png")

    # 📄 Insert both into the .docx file
    doc = Document(doc_path)
    doc.add_paragraph(f"⬆️ Full chart screenshot for {direction} trade")
    doc.add_picture("manual_full.png", width=Inches(5.5))

    doc.add_paragraph(f"🔎 Cropped 4-candle screenshot for {direction} trade")
    doc.add_picture("manual_cropped.png", width=Inches(3.0))
    doc.save(doc_path)

    send_to_clipboard(full_img)
    print(f"✅ Saved both screenshots to {os.path.basename(doc_path)} and copied full chart to clipboard.")

# 🎯 Hotkey guide
print(f"\n📂 Timeframe set to: {timeframe}")
print("📸 Press 'U' to save UP trade")
print("📸 Press 'D' to save DOWN trade")
print("📸 Press 'N' to save NO_SIGNAL trade")
print("❌ Press 'ESC' to exit\n")

keyboard.add_hotkey('u', lambda: capture_and_insert(UP_PATH, "UP"))
keyboard.add_hotkey('d', lambda: capture_and_insert(DOWN_PATH, "DOWN"))
keyboard.add_hotkey('n', lambda: capture_and_insert(NO_SIGNAL_PATH, "NO_SIGNAL"))
keyboard.wait('esc')
print("👋 Exiting.")
